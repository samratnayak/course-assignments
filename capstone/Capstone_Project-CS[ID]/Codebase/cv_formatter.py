"""
CV Formatter Module

This module handles formatting and exporting CVs to DOCX and PDF formats.
"""

import os
import re
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from typing import Dict, Optional, List, Any
import warnings

warnings.filterwarnings("ignore")

# Document.add_paragraph and table cell both support add_paragraph
_AddParagraphParent = Any


def _set_cell_shading(cell, fill_hex: str) -> None:
    """Light background for sidebar cells (e.g. F0F4F7)."""
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill_hex)
    tc_pr.append(shd)


def _set_paragraph_bottom_border(paragraph, color: str = "A0A0A0", size: str = "4") -> None:
    """Draw a single bottom rule on the paragraph (OOXML ``w:pBdr``). Used for classic section headers."""
    p_pr = paragraph._element.get_or_add_pPr()
    p_bdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), size)
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), color)
    p_bdr.append(bottom)
    p_pr.append(p_bdr)


class CVFormatter:
    """
    Class to format and export CVs to DOCX and PDF formats.
    """
    
    def __init__(self):
        """Initialize the CVFormatter."""
        pass
    
    def _remove_quotes_from_text(self, text: str) -> str:
        """
        Remove leading and trailing quotes from text, handling both single and double quotes.
        
        Args:
            text: Text that may contain quotes
            
        Returns:
            Text with quotes removed
        """
        if not text:
            return text
        
        text = text.strip()
        
        # Remove quotes that wrap the entire text
        if text.startswith('"') and text.endswith('"'):
            text = text[1:-1].strip()
        elif text.startswith("'") and text.endswith("'"):
            text = text[1:-1].strip()
        
        # Remove quotes from start/end individually
        text = text.strip('"').strip("'").strip()
        
        return text
    
    def format_to_docx(
        self,
        cv_sections: Dict[str, str],
        output_path: str,
        candidate_name: str = "Candidate",
        document_style: str = "default",
    ):
        """
        Export CV to DOCX. ``document_style`` only changes document layout, not section text.
        """
        doc = Document()
        style_key = (document_style or "default").strip().lower()
        if style_key == "sidebar_modern":
            self._format_to_docx_sidebar_modern(doc, cv_sections, candidate_name)
        elif style_key == "classic_serif":
            self._format_to_docx_classic_serif(doc, cv_sections, candidate_name)
        else:
            self._format_to_docx_default(doc, cv_sections, candidate_name)
        doc.save(output_path)
        print(f"CV saved to: {output_path}")

    def _format_to_docx_default(self, doc: Document, cv_sections: Dict[str, str], candidate_name: str) -> None:
        """Original centered-name, blue-accent single-column layout."""
        # Set compact document margins for 2-3 page CV
        sections = doc.sections
        for section in sections:
            section.top_margin = Inches(0.5)
            section.bottom_margin = Inches(0.5)
            section.left_margin = Inches(0.7)
            section.right_margin = Inches(0.7)
        
        # Modern header with candidate name
        header_para = doc.add_paragraph()
        header_run = header_para.add_run(candidate_name)
        header_run.font.size = Pt(28)
        header_run.font.bold = True
        header_run.font.color.rgb = RGBColor(26, 35, 126)  # Professional dark blue
        header_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        header_para.paragraph_format.space_after = Pt(8)
        
        # Add sections with compact styling for 2-3 page CV
        for section_name, section_content in cv_sections.items():
            # Add section heading with compact styling
            heading_para = doc.add_paragraph()
            heading_run = heading_para.add_run(section_name.upper())
            heading_run.font.size = Pt(12)
            heading_run.font.bold = True
            heading_run.font.color.rgb = RGBColor(26, 35, 126)  # Professional dark blue (matches header)
            heading_para.paragraph_format.space_before = Pt(12)
            heading_para.paragraph_format.space_after = Pt(6)
            
            # Add a compact colored underline/separator
            separator_para = doc.add_paragraph()
            separator_run = separator_para.add_run('_' * 50)  # Clean separator line
            separator_run.font.color.rgb = RGBColor(26, 35, 126)  # Matching blue
            separator_run.font.size = Pt(9)
            separator_para.paragraph_format.space_after = Pt(8)
            separator_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
            
            # Add section content with improved formatting
            # Special handling for Work Experience section
            if section_name.lower() in ["work experience", "experience"]:
                self._format_work_experience_section(doc, section_content)
            else:
                # Clean quotes from the entire section content first
                section_content = self._remove_quotes_from_text(section_content)
                
                # Special cleaning for Personal Information - remove "Contact Information" heading and placeholders
                if section_name == "Personal Information":
                    # Remove "Contact Information" lines and placeholders
                    lines = section_content.split('\n')
                    filtered_lines = []
                    for line in lines:
                        line_stripped = line.strip()
                        if not line_stripped:
                            continue
                        line_lower = line_stripped.lower()
                        # Skip lines that are just "Contact Information" or "Contact Information:"
                        if line_lower in ['contact information', 'contact information:', 'contact information :']:
                            continue
                        # Skip lines that start with "Contact Information"
                        if line_lower.startswith('contact information'):
                            remaining = line_lower.replace('contact information', '').replace(':', '').strip()
                            if len(remaining) < 5:  # Very little remaining text, likely just the heading
                                continue
                        # Remove placeholder text patterns
                        placeholder_patterns = [
                            r'\[linkedin\s+url\s+if\s+provided\]',
                            r'\[address\s+if\s+provided\]',
                            r'\[linkedin\s+url\s+if\s+available\]',
                            r'\[address\s+if\s+available\]',
                            r'linkedin\s+url\s+if\s+provided',
                            r'address\s+if\s+provided',
                            r'linkedin\s+url\s+if\s+available',
                            r'address\s+if\s+available',
                        ]
                        is_placeholder = False
                        for pattern in placeholder_patterns:
                            if re.search(pattern, line_lower):
                                is_placeholder = True
                                break
                        if is_placeholder:
                            continue
                        # Also check for lines that are just brackets with placeholder-like text
                        if line_stripped.startswith('[') and line_stripped.endswith(']'):
                            if any(keyword in line_lower for keyword in ['if provided', 'if available', 'linkedin', 'address']):
                                continue
                        filtered_lines.append(line)
                    section_content = '\n'.join(filtered_lines).strip()

                paragraphs = section_content.split('\n')
                for para_text in paragraphs:
                    para_text = para_text.strip()
                    if not para_text:
                        continue
                    
                    # Remove leading and trailing quotes from each line
                    para_text = self._remove_quotes_from_text(para_text)
                    
                    # Skip "Contact Information" lines and placeholders in Personal Information section
                    if section_name == "Personal Information":
                        para_lower = para_text.lower().strip()
                        if para_lower in ['contact information', 'contact information:', 'contact information :']:
                            continue
                        if para_lower.startswith('contact information') and len(para_lower.replace('contact information', '').replace(':', '').strip()) < 5:
                            continue
                        # Check for placeholder patterns
                        placeholder_patterns = [
                            r'\[linkedin\s+url\s+if\s+provided\]',
                            r'\[address\s+if\s+provided\]',
                            r'\[linkedin\s+url\s+if\s+available\]',
                            r'\[address\s+if\s+available\]',
                        ]
                        for pattern in placeholder_patterns:
                            if re.search(pattern, para_lower):
                                continue
                        # Check for bracket-wrapped placeholders
                        if para_text.strip().startswith('[') and para_text.strip().endswith(']'):
                            if any(keyword in para_lower for keyword in ['if provided', 'if available', 'linkedin', 'address']):
                                continue
                    
                    if not para_text:
                        continue
                    
                    # Skip if it's a section title that wasn't cleaned (additional safety check)
                    section_name_upper = section_name.upper()
                    if (section_name_upper in para_text.upper() and len(para_text) < 50 and 
                        (para_text.upper() == section_name_upper or 
                         para_text.upper().startswith(section_name_upper) or
                         f"**{section_name_upper}**" in para_text.upper())):
                        continue
                    
                    # Check if it's a bullet point
                    if para_text.startswith('-') or para_text.startswith('•') or para_text.startswith('*'):
                        bullet_text = para_text.lstrip('- •*').strip()
                        # Remove quotes from bullet text too
                        bullet_text = bullet_text.strip('"').strip("'").strip()
                        if bullet_text:
                            para = doc.add_paragraph(bullet_text, style='List Bullet')
                            # Customize bullet style with compact spacing
                            para_format = para.paragraph_format
                            para_format.space_after = Pt(4)
                            para_format.left_indent = Inches(0.25)
                            para_format.first_line_indent = Inches(-0.25)
                            
                            # Style bullet text
                            for run in para.runs:
                                run.font.size = Pt(10)
                                run.font.color.rgb = RGBColor(44, 62, 80)  # Dark blue-gray
                    else:
                        para = doc.add_paragraph(para_text)
                        para_format = para.paragraph_format
                        para_format.space_after = Pt(5)
                        para_format.left_indent = Inches(0.1)
                        
                        # Style the paragraph text
                        for run in para.runs:
                            run.font.size = Pt(11)
                            run.font.color.rgb = RGBColor(44, 62, 80)  # Dark blue-gray

                # Add spacing between sections
                doc.add_paragraph()
    
    def _format_work_experience_section(self, parent: _AddParagraphParent, section_content: str):
        """
        Format Work Experience section with bold subheadings for Company, Position, and Dates.
        
        Args:
            parent: Document or table cell (supports add_paragraph)
            section_content: Content of the work experience section
        """
        
        # Remove markdown formatting (asterisks) from the entire content first
        section_content = re.sub(r'\*\*([^*]+)\*\*', r'\1', section_content)  # Remove **bold**
        section_content = re.sub(r'\*([^*]+)\*', r'\1', section_content)  # Remove *italic*
        
        paragraphs = section_content.split('\n')
        current_company = None
        current_position = None
        current_dates = None
        current_bullets = []
        seen_positions = set()  # Track seen positions to avoid duplicates
        
        for para_text in paragraphs:
            para_text = para_text.strip()
            if not para_text:
                continue
            
            # Remove markdown formatting from the line
            para_text = re.sub(r'\*\*([^*]+)\*\*', r'\1', para_text)  # Remove **bold**
            para_text = re.sub(r'\*([^*]+)\*', r'\1', para_text)  # Remove *italic*
            para_text = para_text.strip()
            
            # Skip section title if present
            if any(keyword in para_text.upper() for keyword in ['WORK EXPERIENCE', 'EXPERIENCE']):
                continue
            
            # Check if this line looks like a job title/company/date line
            # Patterns: "Job Title", "Company Name", "Date Range" or "Job Title at Company, Date"
            # Try to detect company names (often contain "Inc", "LLC", "Corp", "Ltd", or are standalone)
            # Try to detect dates (contain months, years, "Present", "-", "to")
            date_pattern = r'\b(Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Oct|Nov|Dec)[a-z]*\s+\d{4}|Present|\d{4}\s*-\s*\d{4}|\d{4}\s*to\s*\d{4}'
            has_date = bool(re.search(date_pattern, para_text, re.IGNORECASE))
            
            # If line has date, it's likely a position/company/date line
            # Also check if it looks like a position/company line (not too long, no bullet markers)
            is_likely_header = (has_date or (len(para_text) < 100 and not para_text.startswith('-') 
                            and not para_text.startswith('•') and not para_text.startswith('*')))
            
            if is_likely_header:
                # Remove markdown first for comparison
                para_text_clean_for_compare = para_text.strip()
                para_text_clean_for_compare = re.sub(r'\*\*([^*]+)\*\*', r'\1', para_text_clean_for_compare)
                para_text_clean_for_compare = re.sub(r'\*([^*]+)\*', r'\1', para_text_clean_for_compare)
                para_text_clean_for_compare = para_text_clean_for_compare.strip().lower()
                
                # Check if this is a duplicate position (same position as current, without markdown)
                if current_position:
                    current_pos_clean = re.sub(r'\*\*([^*]+)\*\*', r'\1', current_position)
                    current_pos_clean = re.sub(r'\*([^*]+)\*', r'\1', current_pos_clean)
                    current_pos_clean = current_pos_clean.strip().lower()
                    
                    # If this line matches the current position (after removing markdown), skip it
                    if para_text_clean_for_compare == current_pos_clean:
                        continue
                    # Also check if it's just the position repeated (common pattern: **Position** then Position)
                    if para_text_clean_for_compare in current_pos_clean or current_pos_clean in para_text_clean_for_compare:
                        if len(para_text_clean_for_compare.split()) <= 3:  # Short line, likely just position
                            continue
                
                # Save previous entry if exists and we're starting a new one
                if (current_company or current_position) and (current_bullets or current_dates):
                    self._add_work_experience_entry(parent, current_company, current_position, current_dates, current_bullets)
                    current_company = None
                    current_position = None
                    current_dates = None
                    current_bullets = []
                
                # Parse the line - could be "Position", "Company, Location", "Date Range"
                # or "Position at Company, Date Range"
                # Remove leading/trailing quotes and markdown first
                para_text = para_text.strip().strip('"').strip("'")
                para_text = re.sub(r'\*\*([^*]+)\*\*', r'\1', para_text)  # Remove **bold**
                para_text = re.sub(r'\*([^*]+)\*', r'\1', para_text)  # Remove *italic*
                para_text = para_text.strip()
                
                parts = [p.strip().strip('"').strip("'") for p in para_text.split(',')]
                # Remove markdown from each part
                parts = [re.sub(r'\*\*([^*]+)\*\*', r'\1', p).strip() for p in parts]
                parts = [re.sub(r'\*([^*]+)\*', r'\1', p).strip() for p in parts]
                
                # Try to identify position, company, and dates
                if len(parts) >= 2:
                    # Could be "Position, Company, Location, Date" or "Position at Company, Date"
                    if ' at ' in parts[0].lower() or ' - ' in parts[0]:
                        # "Position at Company" or "Position - Company"
                        pos_company = parts[0].strip().strip('"').strip("'")
                        pos_company = re.sub(r'\*\*([^*]+)\*\*', r'\1', pos_company)  # Remove markdown
                        if ' at ' in pos_company.lower():
                            pos_parts = pos_company.split(' at ', 1)
                            current_position = pos_parts[0].strip().strip('"').strip("'")
                            current_company = pos_parts[1].strip().strip('"').strip("'")
                        elif ' - ' in pos_company:
                            pos_parts = pos_company.split(' - ', 1)
                            current_position = pos_parts[0].strip().strip('"').strip("'")
                            current_company = pos_parts[1].strip().strip('"').strip("'")
                        else:
                            current_position = pos_company.strip().strip('"').strip("'")
                            current_company = parts[1].strip().strip('"').strip("'") if len(parts) > 1 else None
                    else:
                        # First part might be position, second might be company
                        current_position = parts[0].strip().strip('"').strip("'")
                        current_company = parts[1].strip().strip('"').strip("'") if len(parts) > 1 else None
                    
                    # Last part with date pattern is likely dates
                    for part in reversed(parts):
                        part_cleaned = part.strip().strip('"').strip("'")
                        part_cleaned = re.sub(r'\*\*([^*]+)\*\*', r'\1', part_cleaned)  # Remove markdown
                        if re.search(date_pattern, part_cleaned, re.IGNORECASE):
                            current_dates = part_cleaned
                            break
                else:
                    # Single line - could be position, company, or dates
                    para_text_cleaned = para_text.strip().strip('"').strip("'")
                    para_text_cleaned = re.sub(r'\*\*([^*]+)\*\*', r'\1', para_text_cleaned)  # Remove markdown
                    if has_date:
                        current_dates = para_text_cleaned
                    elif not current_position:
                        # Check if this position was already seen (avoid duplicates)
                        position_key = para_text_cleaned.lower()
                        if position_key not in seen_positions:
                            current_position = para_text_cleaned
                            seen_positions.add(position_key)
                    elif not current_company:
                        current_company = para_text_cleaned
            
            # Bullet points
            elif para_text.startswith('-') or para_text.startswith('•') or para_text.startswith('*'):
                bullet_text = para_text.lstrip('- •*').strip()
                if bullet_text:
                    current_bullets.append(bullet_text)
        
        # Add last entry
        if current_company or current_position:
            self._add_work_experience_entry(parent, current_company, current_position, current_dates, current_bullets)
    
    def _add_work_experience_entry(self, parent: _AddParagraphParent, company: Optional[str], position: Optional[str], 
                                   dates: Optional[str], bullets: List[str]):
        """
        Add a work experience entry with formatted subheadings.
        
        Args:
            parent: Document or table cell
            company: Company name
            position: Job position/title
            dates: Date range
            bullets: List of bullet points
        """
        
        # Clean quotes and markdown from all fields
        if position:
            position = position.strip().strip('"').strip("'")
            position = re.sub(r'\*\*([^*]+)\*\*', r'\1', position)  # Remove **bold**
            position = re.sub(r'\*([^*]+)\*', r'\1', position)  # Remove *italic*
            position = position.strip()
        if company:
            company = company.strip().strip('"').strip("'")
            company = re.sub(r'\*\*([^*]+)\*\*', r'\1', company)  # Remove **bold**
            company = re.sub(r'\*([^*]+)\*', r'\1', company)  # Remove *italic*
            company = company.strip()
        if dates:
            dates = dates.strip().strip('"').strip("'")
            dates = re.sub(r'\*\*([^*]+)\*\*', r'\1', dates)  # Remove **bold**
            dates = re.sub(r'\*([^*]+)\*', r'\1', dates)  # Remove *italic*
            dates = dates.strip()
        
        # Add position (bold) with compact spacing
        if position:
            pos_para = parent.add_paragraph()
            pos_run = pos_para.add_run(position)
            pos_run.font.size = Pt(11)
            pos_run.font.bold = True
            pos_run.font.color.rgb = RGBColor(26, 35, 126)  # Dark blue
            pos_para.paragraph_format.space_after = Pt(2)
        
        # Add company (bold) with compact spacing
        if company:
            comp_para = parent.add_paragraph()
            comp_run = comp_para.add_run(company)
            comp_run.font.size = Pt(10)
            comp_run.font.bold = True
            comp_run.font.color.rgb = RGBColor(52, 73, 94)  # Blue-gray
            comp_para.paragraph_format.space_after = Pt(2)
        
        # Add dates (bold) with compact spacing
        if dates:
            date_para = parent.add_paragraph()
            date_run = date_para.add_run(dates)
            date_run.font.size = Pt(10)
            date_run.font.bold = True
            date_run.font.color.rgb = RGBColor(52, 73, 94)  # Blue-gray
            date_run.font.italic = True
            date_para.paragraph_format.space_after = Pt(3)
        
        # Add bullet points with compact spacing
        for bullet in bullets:
            para = parent.add_paragraph(bullet, style='List Bullet')
            para_format = para.paragraph_format
            para_format.space_after = Pt(3)
            para_format.left_indent = Inches(0.25)
            para_format.first_line_indent = Inches(-0.25)
            
            for run in para.runs:
                run.font.size = Pt(10)
                run.font.color.rgb = RGBColor(44, 62, 80)  # Dark blue-gray
        
        # Add minimal spacing after entry
        spacing_para = parent.add_paragraph()
        spacing_para.paragraph_format.space_after = Pt(4)

    def _sidebar_column_heading(self, cell, title: str) -> None:
        """Uppercase section label plus a short rule (sidebar column)."""
        p = cell.add_paragraph()
        r = p.add_run(title.upper())
        r.font.bold = True
        r.font.size = Pt(9)
        r.font.name = "Calibri"
        r.font.color.rgb = RGBColor(30, 30, 30)
        p.paragraph_format.space_before = Pt(10)
        p.paragraph_format.space_after = Pt(2)
        u = cell.add_paragraph()
        ur = u.add_run("―" * 18)
        ur.font.size = Pt(6)
        ur.font.color.rgb = RGBColor(160, 170, 180)

    def _append_plain_section_body(
        self,
        parent: _AddParagraphParent,
        section_name: str,
        section_content: str,
        body_pt: int = 9,
        font_name: str = "Calibri",
        body_color: RGBColor = None,
    ) -> None:
        """Append plain section lines (not Work Experience). For Personal Information, strips contact headings/placeholders (same rules as default layout)."""
        if body_color is None:
            body_color = RGBColor(40, 55, 70)
        section_content = self._remove_quotes_from_text(section_content)
        if section_name == "Personal Information":
            lines = section_content.split("\n")
            filtered_lines = []
            for line in lines:
                line_stripped = line.strip()
                if not line_stripped:
                    continue
                line_lower = line_stripped.lower()
                if line_lower in ["contact information", "contact information:", "contact information :"]:
                    continue
                if line_lower.startswith("contact information"):
                    remaining = line_lower.replace("contact information", "").replace(":", "").strip()
                    if len(remaining) < 5:
                        continue
                placeholder_patterns = [
                    r"\[linkedin\s+url\s+if\s+provided\]",
                    r"\[address\s+if\s+provided\]",
                    r"\[linkedin\s+url\s+if\s+available\]",
                    r"\[address\s+if\s+available\]",
                ]
                if any(re.search(p, line_lower) for p in placeholder_patterns):
                    continue
                if line_stripped.startswith("[") and line_stripped.endswith("]"):
                    if any(k in line_lower for k in ["if provided", "if available", "linkedin", "address"]):
                        continue
                filtered_lines.append(line)
            section_content = "\n".join(filtered_lines).strip()

        for para_text in section_content.split("\n"):
            para_text = para_text.strip()
            if not para_text:
                continue
            para_text = self._remove_quotes_from_text(para_text)
            if section_name == "Personal Information":
                pl = para_text.lower().strip()
                if pl in ["contact information", "contact information:", "contact information :"]:
                    continue
                if pl.startswith("contact information") and len(pl.replace("contact information", "").replace(":", "").strip()) < 5:
                    continue
            snu = section_name.upper()
            if snu in para_text.upper() and len(para_text) < 50:
                if para_text.upper() == snu or para_text.upper().startswith(snu):
                    continue
            if para_text.startswith("-") or para_text.startswith("•") or para_text.startswith("*"):
                bullet_text = para_text.lstrip("- •*").strip().strip('"').strip("'")
                if bullet_text:
                    para = parent.add_paragraph(bullet_text, style="List Bullet")
                    para.paragraph_format.space_after = Pt(2)
                    para.paragraph_format.left_indent = Inches(0.18)
                    para.paragraph_format.first_line_indent = Inches(-0.18)
                    for run in para.runs:
                        run.font.size = Pt(body_pt)
                        run.font.name = font_name
                        run.font.color.rgb = body_color
            else:
                para = parent.add_paragraph(para_text)
                para.paragraph_format.space_after = Pt(2)
                for run in para.runs:
                    run.font.size = Pt(body_pt)
                    run.font.name = font_name
                    run.font.color.rgb = body_color

    def _format_to_docx_sidebar_modern(
        self, doc: Document, cv_sections: Dict[str, str], candidate_name: str
    ) -> None:
        """Two-column table: shaded left column for contact/education/skills; main body on the right."""
        for section in doc.sections:
            section.top_margin = Inches(0.45)
            section.bottom_margin = Inches(0.45)
            section.left_margin = Inches(0.55)
            section.right_margin = Inches(0.55)

        order = [
            "Personal Information",
            "Professional Summary",
            "Work Experience",
            "Education",
            "Skills",
            "Certifications",
            "Projects",
            "Languages",
        ]
        left_sections = frozenset(
            {"Personal Information", "Education", "Skills", "Certifications", "Languages"}
        )

        table = doc.add_table(rows=1, cols=2)
        table.autofit = False
        cell_l = table.rows[0].cells[0]
        cell_r = table.rows[0].cells[1]
        cell_l.width = Inches(2.35)
        cell_r.width = Inches(4.55)
        _set_cell_shading(cell_l, "F0F4F7")

        hp = cell_l.add_paragraph()
        hr = hp.add_run(candidate_name.strip().upper())
        hr.font.bold = True
        hr.font.size = Pt(15)
        hr.font.name = "Calibri"
        hr.font.color.rgb = RGBColor(26, 54, 93)
        # Professional Summary stays in the right column only (no tagline under the name here;
        # a truncated first line did not fit the narrow sidebar).
        hp.paragraph_format.space_after = Pt(10)

        for name in order:
            if name not in cv_sections or name == "Professional Summary":
                continue
            if name not in left_sections:
                continue
            content = cv_sections[name]
            if name == "Personal Information":
                lines = [ln.strip() for ln in content.split("\n") if ln.strip()]
                c_low = candidate_name.strip().lower()
                rest = [ln for ln in lines if ln.strip().lower() != c_low]
                content = "\n".join(rest) if rest else content
            self._sidebar_column_heading(cell_l, name)
            self._append_plain_section_body(cell_l, name, content, body_pt=8)

        for name in order:
            if name not in cv_sections:
                continue
            if name in left_sections:
                continue
            content = cv_sections[name]
            self._sidebar_column_heading(cell_r, name)
            if name.lower() in ["work experience", "experience"]:
                self._format_work_experience_section(cell_r, content)
            else:
                self._append_plain_section_body(cell_r, name, content, body_pt=10)

        for name, content in cv_sections.items():
            if name in order:
                continue
            self._sidebar_column_heading(cell_r, name)
            if name.lower() in ["work experience", "experience"]:
                self._format_work_experience_section(cell_r, content)
            else:
                self._append_plain_section_body(cell_r, name, content, body_pt=10)

    def _format_to_docx_classic_serif(
        self, doc: Document, cv_sections: Dict[str, str], candidate_name: str
    ) -> None:
        """Single column, Times New Roman, centered header, rules under section titles."""
        for section in doc.sections:
            section.top_margin = Inches(0.65)
            section.bottom_margin = Inches(0.65)
            section.left_margin = Inches(0.85)
            section.right_margin = Inches(0.85)

        title = doc.add_paragraph()
        tr = title.add_run(candidate_name)
        tr.font.name = "Times New Roman"
        tr.font.size = Pt(20)
        tr.font.bold = True
        tr.font.color.rgb = RGBColor(0, 0, 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title.paragraph_format.space_after = Pt(4)

        pi = cv_sections.get("Personal Information", "")
        pi = self._remove_quotes_from_text(pi)
        for line in pi.split("\n"):
            line = line.strip()
            if not line:
                continue
            if line.lower() == candidate_name.strip().lower():
                continue
            cp = doc.add_paragraph()
            cr = cp.add_run(line)
            cr.font.name = "Times New Roman"
            cr.font.size = Pt(10)
            cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
            cp.paragraph_format.space_after = Pt(0)

        doc.add_paragraph()

        order = [
            "Professional Summary",
            "Education",
            "Skills",
            "Work Experience",
            "Projects",
            "Certifications",
            "Languages",
        ]
        seen = set()
        # Contact lines already rendered under the centered name; do not repeat as a section.
        seen.add("Personal Information")
        ink = RGBColor(0, 0, 0)

        def render_block(section_name: str, section_content: str) -> None:
            if section_name == "Personal Information":
                return
            hp = doc.add_paragraph()
            hr = hp.add_run(section_name.upper())
            hr.font.name = "Times New Roman"
            hr.font.size = Pt(11)
            hr.font.bold = True
            hr.font.color.rgb = ink
            hp.paragraph_format.space_before = Pt(10)
            hp.paragraph_format.space_after = Pt(4)
            _set_paragraph_bottom_border(hp, color="B0B0B0", size="6")
            if section_name.lower() in ["work experience", "experience"]:
                self._format_work_experience_classic(doc, section_content)
            else:
                section_content = self._remove_quotes_from_text(section_content)
                for para_text in section_content.split("\n"):
                    para_text = para_text.strip()
                    if not para_text:
                        continue
                    para_text = self._remove_quotes_from_text(para_text)
                    if para_text.startswith("-") or para_text.startswith("•") or para_text.startswith("*"):
                        bt = para_text.lstrip("- •*").strip()
                        if bt:
                            para = doc.add_paragraph(bt, style="List Bullet")
                            para.paragraph_format.space_after = Pt(3)
                            for run in para.runs:
                                run.font.name = "Times New Roman"
                                run.font.size = Pt(11)
                                run.font.color.rgb = ink
                    else:
                        para = doc.add_paragraph(para_text)
                        para.paragraph_format.space_after = Pt(3)
                        for run in para.runs:
                            run.font.name = "Times New Roman"
                            run.font.size = Pt(11)
                            run.font.color.rgb = ink

        for name in order:
            if name in cv_sections:
                render_block(name, cv_sections[name])
                seen.add(name)
        for name, content in cv_sections.items():
            if name not in seen:
                render_block(name, content)

    def _format_work_experience_classic(self, doc: Document, section_content: str) -> None:
        """WE block using Times and dark text (classic layout)."""
        section_content = re.sub(r"\*\*([^*]+)\*\*", r"\1", section_content)
        section_content = re.sub(r"\*([^*]+)\*", r"\1", section_content)
        ink = RGBColor(0, 0, 0)
        navy = RGBColor(26, 54, 93)
        for para_text in section_content.split("\n"):
            para_text = para_text.strip()
            if not para_text:
                continue
            if any(k in para_text.upper() for k in ["WORK EXPERIENCE", "EXPERIENCE"]) and len(para_text) < 40:
                continue
            if para_text.startswith("-") or para_text.startswith("•") or para_text.startswith("*"):
                bt = para_text.lstrip("- •*").strip()
                para = doc.add_paragraph(bt, style="List Bullet")
                para.paragraph_format.space_after = Pt(2)
                for run in para.runs:
                    run.font.name = "Times New Roman"
                    run.font.size = Pt(11)
                    run.font.color.rgb = ink
            else:
                para = doc.add_paragraph()
                run = para.add_run(para_text)
                run.font.name = "Times New Roman"
                run.font.size = Pt(11)
                run.font.bold = True
                run.font.color.rgb = navy
                para.paragraph_format.space_after = Pt(2)
    
    def format_to_text(self, cv_sections: Dict[str, str], output_path: str):
        """
        Format CV sections into a plain text file.
        
        Args:
            cv_sections: Dictionary of CV sections
            output_path: Path to save the text file
        """
        with open(output_path, 'w', encoding='utf-8') as f:
            for section_name, section_content in cv_sections.items():
                f.write(f"\n{'='*80}\n")
                f.write(f"{section_name.upper()}\n")
                f.write(f"{'='*80}\n")
                f.write(f"{section_content}\n\n")
        
        print(f"CV saved to: {output_path}")
    
    def format_to_pdf(
        self,
        cv_sections: Dict[str, str],
        output_path: str,
        candidate_name: str = "Candidate",
        document_style: str = "default",
    ):
        """
        Format CV sections into a PDF document.
        First creates DOCX, then converts to PDF (requires external tool or library).
        
        Args:
            cv_sections: Dictionary of CV sections
            output_path: Path to save the PDF file
            candidate_name: Name of the candidate
        """
        # For now, create DOCX first
        # PDF conversion would require additional libraries like pdfkit or reportlab
        # Or use pandoc if available: pandoc input.docx -o output.pdf
        
        docx_path = output_path.replace('.pdf', '.docx')
        self.format_to_docx(cv_sections, docx_path, candidate_name, document_style=document_style)
        
        # Try to convert to PDF using pandoc if available
        try:
            import subprocess
            result = subprocess.run(
                ['pandoc', docx_path, '-o', output_path],
                capture_output=True,
                timeout=30
            )
            if result.returncode == 0:
                print(f"CV saved to PDF: {output_path}")
                # Optionally remove DOCX file
                # os.remove(docx_path)
            else:
                print(f"PDF conversion failed. DOCX saved at: {docx_path}")
                print("Note: Install pandoc for PDF conversion, or use the DOCX file.")
        except (FileNotFoundError, subprocess.TimeoutExpired):
            print(f"PDF conversion not available. DOCX saved at: {docx_path}")
            print("Note: Install pandoc (https://pandoc.org/installing.html) for PDF conversion.")
    
    def format_cv(
        self,
        cv_sections: Dict[str, str],
        output_path: str,
        format_type: str = "docx",
        candidate_name: str = "Candidate",
        document_style: str = "default",
    ):
        """
        Format CV to specified format.
        
        Args:
            cv_sections: Dictionary of CV sections
            output_path: Path to save the file
            format_type: Output format ('docx', 'txt', 'pdf')
            candidate_name: Name of the candidate
            document_style: Layout preset for docx/pdf only (ignored for txt)
        """
        if format_type.lower() == "docx":
            self.format_to_docx(cv_sections, output_path, candidate_name, document_style=document_style)
        elif format_type.lower() == "txt":
            self.format_to_text(cv_sections, output_path)
        elif format_type.lower() == "pdf":
            self.format_to_pdf(cv_sections, output_path, candidate_name, document_style=document_style)
        else:
            raise ValueError(f"Unsupported format: {format_type}")
