#!/usr/bin/env python3
"""
Script to generate Report.pdf from Report.md
Converts markdown to PDF using available tools.
"""

import os
import sys

def convert_markdown_to_pdf():
    """Convert Report.md to Report.pdf"""
    script_dir = os.path.dirname(os.path.abspath(__file__))
    md_path = os.path.join(script_dir, 'Report.md')
    pdf_path = os.path.join(script_dir, 'Report.pdf')
    
    # Try method 1: Use markdown2pdf or similar
    try:
        import markdown
        from weasyprint import HTML, CSS
        from weasyprint.text.fonts import FontConfiguration
        
        with open(md_path, 'r', encoding='utf-8') as f:
            md_content = f.read()
        
        # Convert markdown to HTML
        html_content = markdown.markdown(md_content, extensions=['tables', 'fenced_code'])
        
        # Add basic styling
        styled_html = f"""
        <!DOCTYPE html>
        <html>
        <head>
            <meta charset="UTF-8">
            <style>
                @page {{
                    size: A4;
                    margin: 1in;
                }}
                body {{
                    font-family: 'Times New Roman', serif;
                    font-size: 11pt;
                    line-height: 1.6;
                }}
                h1 {{
                    font-size: 18pt;
                    font-weight: bold;
                    margin-top: 20pt;
                    margin-bottom: 12pt;
                }}
                h2 {{
                    font-size: 16pt;
                    font-weight: bold;
                    margin-top: 16pt;
                    margin-bottom: 10pt;
                }}
                h3 {{
                    font-size: 14pt;
                    font-weight: bold;
                    margin-top: 12pt;
                    margin-bottom: 8pt;
                }}
                p {{
                    margin-bottom: 6pt;
                }}
                table {{
                    border-collapse: collapse;
                    width: 100%;
                    margin: 10pt 0;
                }}
                th, td {{
                    border: 1px solid #ddd;
                    padding: 8pt;
                    text-align: left;
                }}
                th {{
                    background-color: #f2f2f2;
                    font-weight: bold;
                }}
                ul, ol {{
                    margin-left: 20pt;
                    margin-bottom: 6pt;
                }}
                li {{
                    margin-bottom: 3pt;
                }}
            </style>
        </head>
        <body>
            {html_content}
        </body>
        </html>
        """
        
        # Convert HTML to PDF
        HTML(string=styled_html).write_pdf(pdf_path)
        print(f'✓ Report.pdf created successfully at: {pdf_path}')
        return True
        
    except ImportError:
        pass
    
    # Try method 2: Use markdown2 and pdfkit
    try:
        import markdown2
        import pdfkit
        
        with open(md_path, 'r', encoding='utf-8') as f:
            md_content = f.read()
        
        html_content = markdown2.markdown(md_content, extras=['tables', 'fenced-code-blocks'])
        
        options = {
            'page-size': 'A4',
            'margin-top': '1in',
            'margin-right': '1in',
            'margin-bottom': '1in',
            'margin-left': '1in',
            'encoding': "UTF-8",
            'no-outline': None
        }
        
        pdfkit.from_string(html_content, pdf_path, options=options)
        print(f'✓ Report.pdf created successfully at: {pdf_path}')
        return True
        
    except ImportError:
        pass
    
    # Try method 3: Use pypandoc (wrapper for pandoc)
    try:
        import pypandoc
        
        output = pypandoc.convert_file(
            md_path,
            'pdf',
            format='markdown',
            outputfile=pdf_path,
            extra_args=['--pdf-engine=pdflatex', '-V', 'geometry:margin=1in', '-V', 'fontsize=11pt']
        )
        print(f'✓ Report.pdf created successfully at: {pdf_path}')
        return True
        
    except (ImportError, Exception) as e:
        print(f"pypandoc method failed: {e}")
        pass
    
    # Fallback: Use subprocess to call pandoc directly
    try:
        import subprocess
        
        result = subprocess.run(
            ['pandoc', md_path, '-o', pdf_path, '--pdf-engine=pdflatex', 
             '-V', 'geometry:margin=1in', '-V', 'fontsize=11pt'],
            capture_output=True,
            text=True
        )
        
        if result.returncode == 0:
            print(f'✓ Report.pdf created successfully at: {pdf_path}')
            return True
        else:
            # Try without pdf-engine
            result = subprocess.run(
                ['pandoc', md_path, '-o', pdf_path],
                capture_output=True,
                text=True
            )
            if result.returncode == 0:
                print(f'✓ Report.pdf created successfully at: {pdf_path}')
                return True
            else:
                print(f"Error: {result.stderr}")
    except Exception as e:
        print(f"Subprocess method failed: {e}")
    
    # Try method 4: Convert via DOCX (if python-docx is available)
    try:
        from docx import Document
        from docx.shared import Pt, Inches
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        import subprocess
        
        # First generate DOCX
        with open(md_path, 'r', encoding='utf-8') as f:
            content = f.read()
        
        doc = Document()
        sections = doc.sections
        for section in sections:
            section.top_margin = Inches(1)
            section.bottom_margin = Inches(1)
            section.left_margin = Inches(1)
            section.right_margin = Inches(1)
        
        style = doc.styles['Normal']
        font = style.font
        font.name = 'Calibri'
        font.size = Pt(12)
        
        lines = content.split('\n')
        in_table = False
        table = None
        table_headers = None
        
        for i, line in enumerate(lines):
            line = line.strip()
            if not line:
                if not in_table:
                    doc.add_paragraph()
                continue
            
            if line.startswith('# '):
                title = line[2:].strip()
                para = doc.add_heading(title, level=0)
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                para.paragraph_format.space_after = Pt(12)
                in_table = False
            elif line.startswith('## '):
                heading = line[3:].strip()
                doc.add_heading(heading, level=1)
                para = doc.paragraphs[-1]
                para.paragraph_format.space_before = Pt(12)
                para.paragraph_format.space_after = Pt(6)
                in_table = False
            elif line.startswith('### '):
                subheading = line[4:].strip()
                doc.add_heading(subheading, level=2)
                para = doc.paragraphs[-1]
                para.paragraph_format.space_before = Pt(10)
                para.paragraph_format.space_after = Pt(4)
                in_table = False
            elif '|' in line and line.count('|') >= 2:
                cells = [c.strip() for c in line.split('|') if c.strip() and c.strip() != '---']
                if 'Metric' in line or 'Value' in line or 'Notes' in line:
                    table = doc.add_table(rows=1, cols=len(cells))
                    table.style = 'Light Grid Accent 1'
                    table_headers = cells
                    for j, cell_text in enumerate(cells):
                        cell = table.rows[0].cells[j]
                        cell.text = cell_text
                        cell.paragraphs[0].runs[0].font.bold = True
                        cell.paragraphs[0].runs[0].font.size = Pt(11)
                    in_table = True
                elif in_table and table and len(cells) == len(table_headers):
                    row = table.add_row()
                    for j, cell_text in enumerate(cells):
                        if j < len(row.cells):
                            row.cells[j].text = cell_text
                            row.cells[j].paragraphs[0].runs[0].font.size = Pt(11)
            elif line.startswith('- '):
                text = line[2:].strip()
                if '**' in text:
                    para = doc.add_paragraph()
                    parts = text.split('**')
                    for j, part in enumerate(parts):
                        if j % 2 == 0:
                            para.add_run(part)
                        else:
                            run = para.add_run(part)
                            run.bold = True
                else:
                    para = doc.add_paragraph(text, style='List Bullet')
                para.paragraph_format.space_after = Pt(3)
                in_table = False
            elif line.startswith('**') and line.endswith('**'):
                para = doc.add_paragraph()
                run = para.add_run(line[2:-2])
                run.bold = True
                in_table = False
            else:
                if '**' in line:
                    para = doc.add_paragraph()
                    parts = line.split('**')
                    for j, part in enumerate(parts):
                        if j % 2 == 0:
                            para.add_run(part)
                        else:
                            run = para.add_run(part)
                            run.bold = True
                else:
                    para = doc.add_paragraph(line)
                para.paragraph_format.space_after = Pt(6)
                in_table = False
        
        docx_path = os.path.join(script_dir, 'Report_temp.docx')
        doc.save(docx_path)
        
        # Try to convert DOCX to PDF using LibreOffice or similar
        result = subprocess.run(
            ['soffice', '--headless', '--convert-to', 'pdf', '--outdir', script_dir, docx_path],
            capture_output=True,
            text=True
        )
        
        if result.returncode == 0:
            # Rename the output
            temp_pdf = os.path.join(script_dir, 'Report_temp.pdf')
            if os.path.exists(temp_pdf):
                if os.path.exists(pdf_path):
                    os.remove(pdf_path)
                os.rename(temp_pdf, pdf_path)
                os.remove(docx_path)
                print(f'✓ Report.pdf created successfully at: {pdf_path}')
                return True
        
        # Clean up temp file
        if os.path.exists(docx_path):
            os.remove(docx_path)
            
    except (ImportError, Exception) as e:
        pass
    
    # If all methods fail, provide instructions
    print("\n❌ Could not automatically generate PDF. Please use one of these methods:")
    print("\n**Recommended: Use LibreOffice (if installed)**")
    print("   1. Run: python3 generate_report.py  (creates Report.docx)")
    print("   2. Open Report.docx in LibreOffice")
    print("   3. File → Export as PDF")
    print("\n**Alternative: Install PDF generation tools**")
    print("   pip install markdown weasyprint")
    print("   Then run: python3 generate_pdf.py")
    print("\n**Or use pandoc with LaTeX:**")
    print("   Install: brew install basictex  (macOS)")
    print("   Then: pandoc Report.md -o Report.pdf --pdf-engine=pdflatex")
    
    return False

if __name__ == '__main__':
    success = convert_markdown_to_pdf()
    sys.exit(0 if success else 1)
