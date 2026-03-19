#!/usr/bin/env python3
"""
Script to generate Report.docx from Report.md
Run this script to convert the markdown report to Word format.
"""

try:
    from docx import Document
    from docx.shared import Pt, Inches, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    import os
    
    # Read the markdown content
    script_dir = os.path.dirname(os.path.abspath(__file__))
    md_path = os.path.join(script_dir, 'Report.md')
    
    with open(md_path, 'r', encoding='utf-8') as f:
        content = f.read()
    
    # Create Word document
    doc = Document()
    
    # Set margins (1 inch all around - standard academic format)
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
    
    # Set default font to Times New Roman (academic standard - matches sample)
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(12)  # Standard 12pt as per sample
    style.paragraph_format.line_spacing = 1.0  # Single spacing to match sample
    style.paragraph_format.space_after = Pt(6)
    style.paragraph_format.first_line_indent = Inches(0)  # No first line indent
    
    # Update heading styles to match sample
    heading1_style = doc.styles['Heading 1']
    heading1_font = heading1_style.font
    heading1_font.name = 'Times New Roman'
    heading1_font.size = Pt(12)  # Same as body text for numbered sections
    heading1_font.bold = True
    heading1_style.paragraph_format.space_before = Pt(12)
    heading1_style.paragraph_format.space_after = Pt(6)
    heading1_style.paragraph_format.left_indent = Inches(0)
    
    heading2_style = doc.styles['Heading 2']
    heading2_font = heading2_style.font
    heading2_font.name = 'Times New Roman'
    heading2_font.size = Pt(12)
    heading2_font.bold = True
    heading2_style.paragraph_format.space_before = Pt(10)
    heading2_style.paragraph_format.space_after = Pt(4)
    
    # Parse and add content
    lines = content.split('\n')
    in_table = False
    table = None
    table_headers = None
    
    for i, line in enumerate(lines):
        line = line.strip()
        if not line:
            if not in_table:
                doc.add_paragraph()
            continue
        
        if line.startswith('# '):
            # Main title - centered, bold, larger font (matches sample)
            title = line[2:].strip()
            para = doc.add_paragraph()
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = para.add_run(title)
            run.font.name = 'Times New Roman'
            run.font.size = Pt(14)  # Match sample title size
            run.font.bold = True
            para.paragraph_format.space_before = Pt(0)
            para.paragraph_format.space_after = Pt(12)
            in_table = False
        elif line.startswith('## '):
            # Section heading - check if it's "Abstract" (special case) or numbered section
            heading = line[3:].strip()
            
            # Special handling for "Abstract" - it's not numbered in the sample
            if heading.lower() == 'abstract':
                para = doc.add_paragraph()
                run = para.add_run(heading)
                run.font.name = 'Times New Roman'
                run.font.size = Pt(12)
                run.font.bold = True
                para.paragraph_format.space_before = Pt(12)
                para.paragraph_format.space_after = Pt(6)
            else:
                # Numbered section (e.g., "1. Introduction", "2. Problem Statement")
                para = doc.add_heading(heading, level=1)
                # Ensure Times New Roman and proper size
                for run in para.runs:
                    run.font.name = 'Times New Roman'
                    run.font.size = Pt(12)  # Same as body text for numbered sections
                    run.font.bold = True
                para.paragraph_format.space_before = Pt(12)
                para.paragraph_format.space_after = Pt(6)
            in_table = False
        elif line.startswith('### '):
            # Subsection (e.g., "4.1 Tools and Technologies")
            subheading = line[4:].strip()
            para = doc.add_heading(subheading, level=2)
            # Ensure Times New Roman
            for run in para.runs:
                run.font.name = 'Times New Roman'
                run.font.size = Pt(12)
            para.paragraph_format.space_before = Pt(10)
            para.paragraph_format.space_after = Pt(4)
            in_table = False
        elif '|' in line and line.count('|') >= 2:
            # Skip table separator lines (---)
            if all(c in ['-', '|', ':', ' '] for c in line.replace('|', '').strip()):
                continue
            
            # Table row
            cells = [c.strip() for c in line.split('|') if c.strip()]
            
            # Check if this is a header row (check next line for separator)
            is_header = False
            if i + 1 < len(lines):
                next_line = lines[i + 1].strip()
                if '---' in next_line or all(c in ['-', '|', ':', ' '] for c in next_line.replace('|', '').strip()):
                    is_header = True
            
            if is_header and not in_table:
                # Table header row
                table = doc.add_table(rows=1, cols=len(cells))
                table.style = 'Light Grid Accent 1'
                table_headers = cells
                for j, cell_text in enumerate(cells):
                    cell = table.rows[0].cells[j]
                    cell.text = cell_text
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            run.font.name = 'Times New Roman'
                            run.font.size = Pt(11)
                            run.font.bold = True
                in_table = True
            elif in_table and table and len(cells) == len(table_headers):
                # Table data row
                row = table.add_row()
                for j, cell_text in enumerate(cells):
                    if j < len(row.cells):
                        row.cells[j].text = cell_text
                        for paragraph in row.cells[j].paragraphs:
                            for run in paragraph.runs:
                                run.font.name = 'Times New Roman'
                                run.font.size = Pt(11)
        elif line.startswith('- '):
            # Bullet point - use bullet symbol (•) to match sample
            text = line[2:].strip()
            # Handle bold text in bullets
            if '**' in text:
                para = doc.add_paragraph()
                # Add bullet symbol
                bullet_run = para.add_run('•\t')
                bullet_run.font.name = 'Times New Roman'
                bullet_run.font.size = Pt(12)
                parts = text.split('**')
                for j, part in enumerate(parts):
                    if j % 2 == 0:
                        run = para.add_run(part)
                        run.font.name = 'Times New Roman'
                        run.font.size = Pt(12)
                    else:
                        run = para.add_run(part)
                        run.font.name = 'Times New Roman'
                        run.font.size = Pt(12)
                        run.bold = True
            else:
                para = doc.add_paragraph()
                # Add bullet symbol manually to match sample
                bullet_run = para.add_run('•\t')
                bullet_run.font.name = 'Times New Roman'
                bullet_run.font.size = Pt(12)
                text_run = para.add_run(text)
                text_run.font.name = 'Times New Roman'
                text_run.font.size = Pt(12)
            para.paragraph_format.space_after = Pt(3)
            para.paragraph_format.left_indent = Inches(0.25)
            para.paragraph_format.first_line_indent = Inches(-0.25)  # Hanging indent
            in_table = False
        elif line.startswith('**') and line.endswith('**'):
            # Bold paragraph
            para = doc.add_paragraph()
            run = para.add_run(line[2:-2])
            run.font.name = 'Times New Roman'
            run.font.size = Pt(12)
            run.bold = True
            para.paragraph_format.space_after = Pt(6)
            in_table = False
        else:
            # Regular paragraph
            # Handle bold text
            if '**' in line:
                para = doc.add_paragraph()
                parts = line.split('**')
                for j, part in enumerate(parts):
                    if j % 2 == 0:
                        run = para.add_run(part)
                        run.font.name = 'Times New Roman'
                        run.font.size = Pt(12)
                    else:
                        run = para.add_run(part)
                        run.font.name = 'Times New Roman'
                        run.font.size = Pt(12)
                        run.bold = True
            else:
                para = doc.add_paragraph(line)
                # Ensure Times New Roman for all runs
                for run in para.runs:
                    if run.font.name is None or run.font.name != 'Times New Roman':
                        run.font.name = 'Times New Roman'
                        run.font.size = Pt(12)
            para.paragraph_format.space_after = Pt(6)
            in_table = False
    
    # Save document
    output_path = os.path.join(script_dir, 'Report.docx')
    doc.save(output_path)
    print(f'✓ Report.docx created successfully at: {output_path}')
    print(f'  Total pages: Approximately {len(doc.paragraphs) // 25} pages')
    
except ImportError:
    print("Error: python-docx library not found.")
    print("Please install it using: pip install python-docx")
    print("\nAlternatively, you can:")
    print("1. Open Report.md in a markdown editor")
    print("2. Export/convert it to Word format")
    print("3. Or use pandoc: pandoc Report.md -o Report.docx")
except Exception as e:
    print(f"Error creating Word document: {e}")
    import traceback
    traceback.print_exc()
